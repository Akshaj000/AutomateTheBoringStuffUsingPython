import docx

doc = docx.Document('invitation.docx')
names = open('guests.txt')
names = names.read()
names = names.split(sep='\n')


for i in names:
    doc.add_paragraph('It would be a pleasure to have the company of').style = "styleone"
    #i was lazy to add styles
    doc.add_paragraph(i)
    doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    doc.add_paragraph('April 1st')
    date = doc.add_paragraph('at 7 o\'clock')
    date.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)


doc.save('invitations.docx')